from pathlib import Path

from docx import Document


DOCX = Path("final/Harpreet_Singh_Embedded_AI_Resume.docx")


REPLACEMENTS = {
    "Build applied neuroengineering systems for BCI and brain-care research, including EEG/fNIRS acquisition, embedded device control, firmware, and Python analysis/visualization workflows.": "Build current PhD neuroengineering research stack for EEG/fNIRS acquisition, CGX/XDF dataset creation, live BCI classifier/game control, embedded device control, and FPGA temporal QiSNN acceleration.",
    "Created Sheeg research software that orchestrates CGX EEG, a behavioural game, UDP-to-LSL event markers, LabRecorder control, XDF output, manifests, stream probes, and session summaries.": "Built CGX/Sheeg workflows that launch CGX Acquisition and LabRecorder, record labelled visual-stimulus XDF sessions, preprocess EEG into CSV datasets, and export model artifacts for live classification.",
    "Built FPGA research pipelines for ANN, SNN, QSNN/QiSNN, QLIF, PC-DDM-SNN, and rodent-raster decision experiments using Vitis HLS, Vivado, MicroBlaze, BRAM windows, and UART diagnostics.": "Developed closed-loop BCI/FPGA and rodent-decision pipelines using ADS1299 frames, MicroBlaze, BRAM windows, temporal QiSNN/SNN accelerators, UDP/UART intent packets, and five-state game control.",
    "snn_qisnn_rodent_raster | Rodent neural raster to temporal QiSNN/SNN FPGA pipeline": "rodent_decision_qisnn_temporal and BCI_Game_Loop_FPGA | Temporal QiSNN decision pipelines",
    "Prepared rodent decision decoding pipeline that converts neural population raster windows into 12 x 196 temporal features for no-lick, left-lick, and right-lick prediction experiments.": "Prepared a temporal QiSNN/SNN FPGA research pipeline for rodent decision prediction using DANDI-derived raster vectors, 12 time bins x 196 features, 2352-word Q10 input BRAM, and 3-class lick labels.",
    "Adapted the active temporal QiSNN hardware interface with 2352-word input BRAM, AXI4-Lite ap_ctrl_hs control, fixed-point payloads, MicroBlaze readback, and debug-visible state BRAM planning.": "Designed the FPGA/electronics branch of a closed-loop BCI stack where ADS1299 frames feed MicroBlaze input BRAM, an EEG-adapted active temporal QiSNN produces five intent scores, and Ethernet/Wi-Fi/UART/SPI transports drive a PC/phone game.",
    "Documented active QiSNN fixed-point evaluation using deployed HLS weights with 96.34% MNIST14 accuracy plus robustness sweep outputs for noise, dropout, salt-pepper, and quantization conditions.": "Connected trained/quantized EEG intent models to a shared software/hardware contract with idle, left, right, up, and down classes, Q15 confidence scoring, golden-window replay, and safety gates before live commands.",
    "8eeg8fNIRS and Sheeg | Neurophysiology acquisition and synchronized behavioural experiment software": "CGX_dataset_game, BCT_8EEG_8FNIRS, and Sheeg | Closed-loop BCI research stack",
    "Integrated ADS1299 EEG, ADS8688 fNIRS analog sampling, DAC8565 output control, watchdog PWM, TDM MUX sequencing, and Raspberry Pi SPI/GPIO concurrency into live acquisition scripts.": "Integrated ADS1299 EEG, ADS8688 fNIRS analog sampling, DAC8565 output control, watchdog PWM, TDM MUX sequencing, shared SPI0 locking, and Raspberry Pi GPIO concurrency into live acquisition scripts.",
    "Built Sheeg session orchestration for EEG + behavioural game recording with CGX Acquisition, UDP-to-LSL marker bridge, stream probing, LabRecorder remote control, XDF recording, and session metadata.": "Built CGX EEG dataset and game workflow: five-state visual-stimulus labels are recorded to XDF, preprocessed into labelled CSV, trained into classifier artifacts, and replayed live as INTENT UDP packets to a browser game.",
}


SUMMARY_OLD = (
    "Embedded AI and neurotechnology researcher with 13+ years across electronics design, firmware, FPGA acceleration, data engineering, cybersecurity AI, and biomedical research. "
    "Current Research Assistant and PhD student at the University of Lethbridge, building EEG/fNIRS acquisition systems, LSL-synchronized research software, and neuromorphic FPGA pipelines for computational behavioural neuroscience. "
    "Strong record translating research ideas into working hardware/software systems, from schematics and PCB bring-up to Vitis HLS, Vivado, MicroBlaze/Zynq integration, Python analytics, and board-visible validation."
)

SUMMARY_NEW = (
    "Embedded AI and neurotechnology researcher with 13+ years across electronics design, firmware, FPGA acceleration, data engineering, cybersecurity AI, and biomedical research. "
    "Current Research Assistant and PhD student at the University of Lethbridge, building EEG/fNIRS acquisition, CGX/XDF dataset creation, live BCI classifier/game control, and temporal QiSNN FPGA pipelines for computational behavioural neuroscience. "
    "Strong record translating research ideas into working hardware/software systems, from schematics and PCB bring-up to Vitis HLS, Vivado, MicroBlaze/Zynq integration, Python analytics, and board-visible validation."
)


def replace_paragraph_text(paragraph, new_text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
    else:
        paragraph.add_run(new_text)


def main():
    doc = Document(DOCX)
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == SUMMARY_OLD:
            replace_paragraph_text(paragraph, SUMMARY_NEW)
            changed += 1
        elif text in REPLACEMENTS:
            replace_paragraph_text(paragraph, REPLACEMENTS[text])
            changed += 1
    if changed < 10:
        raise SystemExit(f"Expected at least 10 replacements, made {changed}")
    doc.save(DOCX)
    print(f"Updated {DOCX} with {changed} research replacements")


if __name__ == "__main__":
    main()
