from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOCX = Path("final/Harpreet_Singh_Embedded_AI_Resume.docx")
SECRYPT = (
    'Harpreet Singh et al., "An Uncertain Reasoning-Based Intrusion Detection '
    'System for DoS/DDoS Detection," SECRYPT 2024, ISBN 978-989-758-709-2, '
    "ISSN 2184-7711, pp. 771-776."
)
THESIS = (
    "Published thesis: Singh H. A robust intrusion detection system utilizing "
    "uncertain reasoning techniques in artificial intelligence. The University "
    "of Regina (Canada); 2024."
)


def insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    if style is not None:
        inserted.style = style
    inserted.add_run(text)
    return inserted


def main():
    doc = Document(DOCX)
    if any(p.text.strip() == THESIS for p in doc.paragraphs):
        print("Thesis citation already present")
        return

    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == SECRYPT:
            insert_paragraph_after(paragraph, THESIS, paragraph.style)
            doc.save(DOCX)
            print(f"Added thesis citation to {DOCX}")
            return

    raise SystemExit("Could not find SECRYPT publication anchor")


if __name__ == "__main__":
    main()
