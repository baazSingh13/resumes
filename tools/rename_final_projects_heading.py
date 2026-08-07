from pathlib import Path

from docx import Document


DOCX = Path("final/Harpreet_Singh_Embedded_AI_Resume.docx")
OLD_HEADING = "Selected GitHub Research & Engineering Projects"
NEW_HEADING = "Research & Engineering Projects"


def main():
    doc = Document(DOCX)
    changed = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == OLD_HEADING:
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = NEW_HEADING
            else:
                paragraph.add_run(NEW_HEADING)
            changed += 1

    if changed != 1:
        raise SystemExit(f"Expected 1 heading replacement, made {changed}")

    doc.save(DOCX)
    print(f"Renamed heading in {DOCX}")


if __name__ == "__main__":
    main()
