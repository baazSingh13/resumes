from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("ats")

CONTACT = (
    "Lethbridge, AB | +1 778-536-4298 | workwithharpreetsingh@gmail.com | "
    "linkedin.com/in/harpreet-singh-b42942213/ | github.com/baazSingh13"
)


BASE_EXPERIENCE = [
    {
        "title": "Research Assistant, Core-Hub Neuroengineering Solutions",
        "org": "University of Lethbridge",
        "dates": "Apr 2025 - Present",
        "bullets": [
            "Build current PhD neuroengineering research stack for EEG/fNIRS acquisition, CGX/XDF dataset creation, live BCI classifier/game control, embedded device control, and FPGA temporal QiSNN acceleration.",
            "Developed an 8-channel EEG + 8-channel fNIRS Raspberry Pi platform using ADS1299, ADS8688, DAC8565, SPI/GPIO, TDM MUX sequencing, live plotting, and CSV logging.",
            "Built CGX/Sheeg workflows that launch CGX Acquisition and LabRecorder, record labelled visual-stimulus XDF sessions, preprocess EEG into CSV datasets, and export model artifacts for live classification.",
            "Developed closed-loop BCI/FPGA and rodent-decision pipelines using ADS1299 frames, MicroBlaze, BRAM windows, temporal QiSNN/SNN accelerators, UDP/UART intent packets, and five-state game control.",
        ],
    },
    {
        "title": "Research Assistant",
        "org": "University of Regina",
        "dates": "Jan 2021 - Apr 2024",
        "bullets": [
            "Researched uncertain-reasoning intrusion detection for DoS/DDoS detection using Bayesian Networks, Markov Networks, Zeek, Wireshark, feature engineering, and Python ML workflows.",
            "Published SECRYPT 2024 work on uncertain-reasoning IDS methods and presented findings to an international cybersecurity research audience.",
            "Built ML prototypes including a Bayesian liver-disorder diagnostic model with 85% accuracy and Punjabi BERT/NLP classifiers for 100K+ news articles with 92% classification accuracy.",
        ],
    },
    {
        "title": "Independent Embedded Systems Developer",
        "org": "Self-Employed",
        "dates": "Jun 2024 - Mar 2025",
        "bullets": [
            "Designed a custom STM32F405RGT6 MicroPython development board, including multilayer PCB design, firmware flashing, sensor/actuator interfaces, power management, and hardware/software debugging.",
        ],
    },
    {
        "title": "Senior Software Engineer",
        "org": "Planetcast Media Services Ltd.",
        "dates": "Dec 2018 - Nov 2020",
        "bullets": [
            "Led a 15-member team delivering embedded hardware and firmware for broadcast/media systems, including a network-controlled 10-channel video switcher and Ethernet temperature acquisition platform.",
            "Designed schematics and PCBs, selected components, collaborated on cabinet/mechanical integration, and developed C firmware with Ethernet, HTTP/LWIP, alarm, and control interfaces.",
        ],
    },
    {
        "title": "Embedded Software Engineer / R&D Engineer",
        "org": "Spark Eighteen Pvt. Ltd.; Exicom Tele-systems Ltd.",
        "dates": "Feb 2016 - Dec 2018",
        "bullets": [
            "Developed STM32/AVR embedded systems for LiFi control, telecom power plants, DC interface cards, and automatic test equipment.",
            "Designed 15-20 W isolated SMPS sections with 78-81% measured efficiency and performed board bring-up across CAN, SPI Flash, I2C EEPROM, Ethernet/LWIP, and GUI test automation.",
        ],
    },
]


PROJECTS = {
    "fpga": [
        {
            "name": "BCI_Game_Loop_FPGA",
            "desc": "Designed the FPGA/electronics branch of a closed-loop BCI stack on Arty A7-100T: ADS1299 frames feed MicroBlaze input BRAM, an EEG-adapted active temporal QiSNN produces five intent scores, and Ethernet/Wi-Fi/UART/SPI transports drive a PC/phone game with an idle reject state.",
        },
        {
            "name": "QSNN_EarlyExit_FPGA_Research",
            "desc": "Built an Arty A7-35T QSNN MicroBlaze system with AXI4-Lite control, BRAM input/output buffers, UART diagnostics, and MNIST board tests; recorded 0.698-0.699 ms HLS latency at 100 MHz, 9.87% LUTs, 5.13% FFs, 41.0% BRAM tiles, 3.33% DSPs, zero routing errors, and positive routed timing slack.",
        },
        {
            "name": "rodent_decision_qisnn_temporal",
            "desc": "Prepared a temporal QiSNN/SNN FPGA research pipeline for rodent decision prediction using DANDI-derived raster vectors, 12 time bins x 196 features, 2352-word Q10 input BRAM, 3-class lick labels, AXI4-Lite ap_ctrl_hs control, and debug-visible accelerator state planning.",
        },
    ],
    "neuro": [
        {
            "name": "CGX_dataset_game",
            "desc": "Built the CGX software branch for labelled EEG dataset creation and software evaluation: CGX EEG plus five-state stimulus markers are recorded to XDF, preprocessed into labelled CSV, trained into classifier artifacts, and replayed live as INTENT UDP packets to a browser game.",
        },
        {
            "name": "BCT_8EEG_8FNIRS",
            "desc": "Integrated ADS1299 EEG, ADS8688 fNIRS analog acquisition, DAC8565 output control, watchdog PWM, TDM MUX sequencing, and Raspberry Pi SPI/GPIO concurrency into live acquisition scripts.",
        },
        {
            "name": "BCI_Game_Loop_FPGA",
            "desc": "Connected trained/quantized EEG intent models to an FPGA/game control contract with fixed idle, left, right, up, and down classes, confidence_q15 scoring, golden-window replay, and safety gates before enabling live commands.",
        },
    ],
    "embedded": [
        {
            "name": "BCT_8EEG_8FNIRS",
            "desc": "Developed Raspberry Pi acquisition/control software for ADS1299, ADS8688, DAC8565, shared SPI0, GPIO MUX/TDM sequencing, watchdog PWM, thread-safe SPI transfers, live plotting, and CSV logging.",
        },
        {
            "name": "BCI_Game_Loop_FPGA",
            "desc": "Defined embedded packet and transport contracts for FPGA-to-game control across Ethernet, Wi-Fi, UART, and SPI, including sequence numbers, class IDs, Q15 confidence values, and optional five-score payloads.",
        },
        {
            "name": "PyBoard / STM32 MicroPython Board",
            "desc": "Designed a custom STM32F405RGT6 development board with multilayer PCB, MicroPython firmware, sensor/actuator interfaces, and power-management debugging.",
        },
    ],
    "research": [
        {
            "name": "Current PhD BCI Research Stack",
            "desc": "Built an end-to-end research workflow spanning CGX EEG acquisition, five-state visual-stimulus labels, LabRecorder/XDF recording, preprocessing manifests, held-session model evaluation, quantized model export, live classifier replay, and FPGA/game-loop integration.",
        },
        {
            "name": "QSNN / SNN / ANN FPGA Research",
            "desc": "Defined hardware-aware research methodology comparing ANN, SNN, fixed-step QSNN, and adaptive early-exit QSNN on accuracy, latency, power/energy, FPGA utilization, and robustness.",
        },
        {
            "name": "Rodent Decision Temporal QiSNN Pipeline",
            "desc": "Transformed neural population raster dynamics into temporal fixed-point features for FPGA-oriented no-lick, left-lick, and right-lick decision prediction, with DANDI sample support and documented raster-to-time-series conversion.",
        },
    ],
    "security": [
        {
            "name": "Uncertain Reasoning IDS",
            "desc": "Developed DoS/DDoS intrusion detection models using Bayesian Networks, Markov Networks, feature engineering, Zeek/Wireshark traffic analysis, TensorFlow, PyTorch, and scikit-learn.",
        },
        {
            "name": "IoT / Embedded Security Background",
            "desc": "Combined embedded firmware, networking, secure-device awareness, Ethernet/MQTT/HTTP protocols, and traffic analysis experience for IoT and cyber-physical systems.",
        },
        {
            "name": "AI/ML Research Prototypes",
            "desc": "Built probabilistic medical diagnosis and NLP classification systems, including a Bayesian liver-disorder model with 85% accuracy and Punjabi BERT classifiers with 92% accuracy.",
        },
    ],
}


RESUMES = [
    {
        "file": "ATS_Embedded_Firmware_Resume.docx",
        "title": "Embedded Firmware Engineer",
        "summary": "Senior embedded systems and firmware engineer with 13+ years building microcontroller products, board-support software, PCB bring-up tools, acquisition systems, industrial interfaces, and embedded Linux/Raspberry Pi applications. Strong C, Embedded C, Python, STM32/AVR/PIC/ESP32, SPI, I2C, UART, CAN, Ethernet/LWIP, SMPS, hardware debugging, and sensor-acquisition background, strengthened by current EEG/fNIRS and closed-loop BCI research work.",
        "skills": [
            "C, Embedded C, Python, Bash, bare-metal firmware, FreeRTOS, Zephyr RTOS, Embedded Linux",
            "ARM Cortex, STM32, AVR, PIC, 8051, ESP32, Atmega328P, Atmega2560A, STM32F4",
            "SPI, I2C, UART, CAN, RS-485, Modbus, USB, Ethernet, MQTT, HTTP, WebSocket, LWIP",
            "PCB design, schematic design, board bring-up, JTAG, SWD, oscilloscopes, logic analyzers, SMPS, power management",
            "ADS1299, ADS8688, DAC8565, Raspberry Pi GPIO/SPI, live plotting, CSV logging, sensor validation",
        ],
        "experience_focus": "embedded",
        "projects": PROJECTS["embedded"],
    },
    {
        "file": "ATS_FPGA_HLS_Resume.docx",
        "title": "FPGA / HLS / Hardware Acceleration Engineer",
        "summary": "FPGA and embedded AI engineer focused on Vitis HLS, Vivado, MicroBlaze, Zynq, AXI4-Lite, AXI HP, BRAM data paths, fixed-point neural accelerators, and UART-visible board validation. Built research pipelines for QSNN, QiSNN, SNN, QLIF, PC-DDM-SNN, ANN baselines, N-MNIST, MNIST14, EEG intent decoding, rodent-raster decision decoding, and Zynq-based LLM acceleration.",
        "skills": [
            "Xilinx Vivado, Vitis HLS, HLS C/C++, MicroBlaze, Zynq Cortex-A9, AXI4-Lite, AXI HP, BRAM maps, UARTLite",
            "Arty A7-35T, Arty A7-100T, Arty Z7-20, Artix-7, Zynq-7000, XSA, bitstream, Vitis platform, board bring-up",
            "Fixed-point arithmetic, ap_fixed<16,6>, ap_fixed<12,4>, quantized weights, INT8/INT4 matrix-vector acceleration",
            "ANN, SNN, QSNN, QiSNN, QLIF, PC-DDM-SNN, temporal N-MNIST, MNIST14, rodent raster, early-exit inference",
            "Python, NumPy, PyTorch, TensorFlow/Keras, model export, HLS testbenches, UART test automation, robustness evaluation",
        ],
        "experience_focus": "fpga",
        "projects": PROJECTS["fpga"],
    },
    {
        "file": "ATS_Neurotechnology_BCI_Resume.docx",
        "title": "Neurotechnology / BCI Engineer",
        "summary": "Neurotechnology and BCI research engineer combining computational behavioural neuroscience, embedded systems, EEG/fNIRS acquisition, synchronized behavioural experiments, CGX/XDF dataset creation, Python analysis, live classifier/game loops, LSL/LabRecorder workflows, and FPGA-based neuromorphic research.",
        "skills": [
            "EEG, fNIRS, BCI, CGX EEG, computational behavioural neuroscience, neural data acquisition, behavioural task markers, XDF sessions",
            "ADS1299, ADS8688, DAC8565, Raspberry Pi, SPI, GPIO, TDM MUX, watchdog PWM, live plotting, CSV logging",
            "Lab Streaming Layer, LSL markers, LabRecorder, CGX EEG, UDP marker bridge, stream probing, session manifests",
            "Python, NumPy, Pandas, Matplotlib, scikit-learn, PyTorch, TensorFlow/Keras, biomedical signal analysis",
            "SNN, QSNN, QiSNN, rodent raster, neural population dynamics, FPGA acceleration, MicroBlaze, BRAM, UART diagnostics",
        ],
        "experience_focus": "neuro",
        "projects": PROJECTS["neuro"],
    },
    {
        "file": "ATS_Research_Engineer_Resume.docx",
        "title": "Research Engineer / Applied Scientist",
        "summary": "Research engineer and PhD student with 13+ years spanning embedded systems, applied AI, computational neuroscience, cybersecurity research, FPGA acceleration, and biomedical data acquisition. Experienced at turning research questions into reproducible software/hardware pipelines, labelled datasets, validation plans, experiment artifacts, metrics, publications, and working prototypes across unfamiliar technical domains.",
        "skills": [
            "Research engineering, experimental design, reproducible workflows, validation plans, metrics, documentation, technical writing",
            "Python, NumPy, Pandas, scikit-learn, TensorFlow, PyTorch, Matplotlib, Bayesian Networks, Markov Networks, NLP transformers",
            "Computational neuroscience, EEG/fNIRS, BCI, rodent raster dynamics, SNN, QSNN, QiSNN, adaptive early-exit inference",
            "Embedded systems, Raspberry Pi, STM32, C, firmware, data acquisition, sensors, SPI, I2C, UART, CAN, Ethernet",
            "FPGA acceleration, Vitis HLS, Vivado, MicroBlaze, Zynq, BRAM, AXI4-Lite, UART diagnostics, resource/timing evidence",
        ],
        "experience_focus": "research",
        "projects": PROJECTS["research"],
    },
    {
        "file": "ATS_AI_Cybersecurity_Research_Resume.docx",
        "title": "AI / Cybersecurity Researcher",
        "summary": "AI and cybersecurity researcher with M.Sc. thesis work in uncertain-reasoning intrusion detection, DoS/DDoS detection, network traffic analysis, Bayesian Networks, Markov Networks, Zeek, Wireshark, Python ML, and applied data science. Broader engineering background in embedded systems, IoT, firmware, and FPGA acceleration supports security work on cyber-physical and edge-device systems.",
        "skills": [
            "Intrusion detection systems, DoS/DDoS detection, uncertain reasoning, Bayesian Networks, Markov Networks, probabilistic graphical models",
            "Zeek, Wireshark, TCP/IP, UDP, ICMP, MQTT, HTTP, WebSocket, SSL/TLS concepts, IoT security, embedded device security",
            "Python, TensorFlow, PyTorch, scikit-learn, Pandas, NumPy, feature engineering, Grid Search, K-Fold Cross Validation",
            "Data engineering, data preprocessing, model evaluation, predictive analytics, NLP transformers, Punjabi BERT, medical AI prototypes",
            "Embedded systems, C, firmware, ARM Cortex, STM32, Ethernet/LWIP, sensors, edge AI, research publication and presentation",
        ],
        "experience_focus": "security",
        "projects": PROJECTS["security"],
    },
]


def set_run_font(run, size=10.0, bold=False, italic=False, color="111111"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def setup_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(9.8)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in [
        ("Heading 1", 15, "000000", 8, 4),
        ("Heading 2", 12, "000000", 7, 3),
        ("Heading 3", 10.5, "000000", 4, 2),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    bullet.font.size = Pt(9.2)
    bullet.paragraph_format.left_indent = Inches(0.32)
    bullet.paragraph_format.first_line_indent = Inches(-0.16)
    bullet.paragraph_format.space_after = Pt(2)
    bullet.paragraph_format.line_spacing = 1.04


def para(doc, text="", size=9.8, bold=False, italic=False, align=None, after=4, before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.04
    p.add_run(text)
    return p


def add_role(doc, role, focus):
    p = para(doc, after=1, before=2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(role["title"])
    set_run_font(r, size=10.1, bold=True)
    r = p.add_run(f" | {role['org']} | {role['dates']}")
    set_run_font(r, size=9.3, italic=True, color="444444")

    bullets = role["bullets"]
    if focus == "embedded" and role["title"].startswith("Research Assistant, Core-Hub"):
        bullets = [
            "Build embedded neuroengineering systems for EEG/fNIRS acquisition, SPI/GPIO device control, firmware validation, live plotting, and Python analysis workflows.",
            "Developed Raspberry Pi acquisition/control software using ADS1299, ADS8688, DAC8565, watchdog PWM, TDM MUX sequencing, threading locks, CSV logging, and hardware debug scripts.",
            "Integrated research-session software with CGX EEG, UDP-to-LSL markers, LabRecorder control, stream probing, XDF output, and session metadata.",
        ]
    elif focus == "fpga" and role["title"].startswith("Research Assistant, Core-Hub"):
        bullets = [
            "Build FPGA research pipelines for ANN, SNN, QSNN/QiSNN, QLIF, PC-DDM-SNN, and rodent-raster decision experiments using Vitis HLS, Vivado, MicroBlaze, BRAM windows, and UART diagnostics.",
            "Developed dataset-to-BRAM flows using 196-feature MNIST14 inputs and 2352-word temporal N-MNIST/rodent-raster inputs with fixed-point payloads and software reference checks.",
            "Documented HLS/Vivado/Vitis rebuild flows, XSA/bitstream paths, UART tests, timing/resource evidence, and hardware-claim boundaries.",
        ]
    elif focus == "neuro" and role["title"].startswith("Research Assistant, Core-Hub"):
        bullets = [
            "Build BCI and neurotechnology research systems for EEG/fNIRS acquisition, behavioural task synchronization, embedded device control, and neural data analysis.",
            "Developed 8-channel EEG + 8-channel fNIRS Raspberry Pi acquisition software using ADS1299, ADS8688, DAC8565, TDM MUX sequencing, live plotting, and offline CSV analysis.",
            "Created Sheeg software for CGX EEG, behavioural game execution, UDP-to-LSL event markers, LabRecorder control, XDF recording, stream probes, manifests, and session summaries.",
        ]
    elif focus == "security" and role["title"] == "Research Assistant":
        bullets = [
            "Researched uncertain-reasoning IDS methods for DoS/DDoS detection using Bayesian Networks, Markov Networks, Zeek, Wireshark, feature engineering, and Python ML workflows.",
            "Published SECRYPT 2024 paper on uncertain-reasoning intrusion detection and validated IDS behavior through network traffic preprocessing, model tuning, and experimental evaluation.",
            "Built related ML prototypes including a Bayesian liver-disorder diagnostic model with 85% accuracy and Punjabi BERT/NLP classifiers for 100K+ news articles with 92% accuracy.",
        ]
    for item in bullets:
        add_bullet(doc, item)


def build_one(config):
    doc = Document()
    setup_doc(doc)

    para(doc, "Harpreet Singh", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    para(doc, config["title"], size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=1)
    para(doc, CONTACT, size=8.7, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)

    heading(doc, "Summary")
    para(doc, config["summary"], size=9.5, after=5)

    heading(doc, "Core Skills")
    for item in config["skills"]:
        add_bullet(doc, item)

    heading(doc, "Professional Experience")
    for role in BASE_EXPERIENCE:
        add_role(doc, role, config["experience_focus"])

    heading(doc, "Selected Projects")
    for proj in config["projects"]:
        p = para(doc, after=1, before=2)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(proj["name"])
        set_run_font(r, size=9.7, bold=True)
        add_bullet(doc, proj["desc"])

    heading(doc, "Publication")
    add_bullet(
        doc,
        'Harpreet Singh et al., "An Uncertain Reasoning-Based Intrusion Detection System for DoS/DDoS Detection," SECRYPT 2024, ISBN 978-989-758-709-2, ISSN 2184-7711, pp. 771-776.',
    )

    heading(doc, "Education")
    add_bullet(doc, "PhD in Computational Behavioural Neuroscience, University of Lethbridge - In progress")
    add_bullet(doc, "M.Sc. Computer Science, University of Regina - 2024")
    add_bullet(doc, "Bachelor's equivalent in Electronics & Telecommunication Engineering, AMIETE/IETE - 2012")

    out = OUT_DIR / config["file"]
    doc.save(out)
    print(out)


def main():
    OUT_DIR.mkdir(exist_ok=True)
    for config in RESUMES:
        build_one(config)


if __name__ == "__main__":
    main()
